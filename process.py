# import re
# import os
# from typing import List, Optional

# class DocumentSegmenter:
#     """文档内容分割器，支持按指定分隔符分割文档内容"""
    
#     def __init__(self, delimiter: str = '@@', keep_delimiter: bool = False):
#         """
#         初始化文档分割器
        
#         Args:
#             delimiter: 用于分割文档的分隔符，默认为 '@@'
#             keep_delimiter: 是否在分割后的内容中保留分隔符，默认为 False
#         """
#         self.delimiter = delimiter
#         self.keep_delimiter = keep_delimiter
    
#     def read_file(self, file_path: str) -> str:
#         """
#         读取文档内容
        
#         Args:
#             file_path: 文档文件路径
        
#         Returns:
#             文档内容字符串
#         """
#         if not os.path.exists(file_path):
#             raise FileNotFoundError(f"文件不存在: {file_path}")
        
#         # 根据文件扩展名选择读取方式
#         ext = os.path.splitext(file_path)[1].lower()
        
#         if ext == '.txt':
#             return self._read_text_file(file_path)
#         elif ext == '.docx':
#             return self._read_docx_file(file_path)
#         elif ext == '.pdf':
#             return self._read_pdf_file(file_path)
#         else:
#             raise ValueError(f"不支持的文件格式: {ext}")
    
#     def _read_text_file(self, file_path: str) -> str:
#         """读取文本文件"""
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
    
#     def _read_docx_file(self, file_path: str) -> str:
#         """读取docx文件"""
#         try:
#             from docx import Document
#         except ImportError:
#             raise ImportError("读取docx文件需要安装python-docx库: pip install python-docx")
        
#         doc = Document(file_path)
#         return '\n'.join([para.text for para in doc.paragraphs])
    
#     def _read_pdf_file(self, file_path: str) -> str:
#         """读取PDF文件"""
#         try:
#             import pdfplumber
#         except ImportError:
#             raise ImportError("读取PDF文件需要安装pdfplumber库: pip install pdfplumber")
        
#         with pdfplumber.open(file_path) as pdf:
#             return '\n'.join([page.extract_text() or '' for page in pdf.pages])
    
#     def split_content(self, content: str) -> List[str]:
#         """
#         按指定分隔符分割文档内容
        
#         Args:
#             content: 文档内容字符串
        
#         Returns:
#             分割后的内容列表
#         """
#         if not self.keep_delimiter:
#             # 不保留分隔符，直接分割
#             return [section.strip() for section in content.split(self.delimiter) if section.strip()]
        
#         # 保留分隔符，使用正则表达式分割
#         pattern = re.compile(f'({re.escape(self.delimiter)})')
#         parts = pattern.split(content)
        
#         segments = []
#         current_segment = ''
        
#         for part in parts:
#             if part == self.delimiter:
#                 if current_segment.strip():
#                     segments.append(current_segment.strip())
#                 current_segment = self.delimiter
#             else:
#                 current_segment += part
        
#         if current_segment.strip():
#             segments.append(current_segment.strip())
        
#         return segments
    
#     def process_file(self, file_path: str) -> List[str]:
#         """
#         读取文件并按分隔符分割内容
        
#         Args:
#             file_path: 文档文件路径
        
#         Returns:
#             分割后的内容列表
#         """
#         content = self.read_file(file_path)
#         return self.split_content(content)

# # 示例用法
# if __name__ == "__main__":
#     segmenter = DocumentSegmenter(delimiter='@@', keep_delimiter=False)
    
#     # 处理文档文件
#     file_path = '/bigdata/lihuanjia/ocr/data/ERDOS_零售管理项目_培训手册_WEBPOS_20250210V1.docx'  # 替换为实际文件路径
#     segments = segmenter.process_file(file_path)
    
#     # 输出分割结果
#     print(f"成功从文档中分割出 {len(segments)} 个内容块")
#     for i, segment in enumerate(segments, 1):
#         print(f"\n=== 第 {i} 个内容块 ===")
#         print(segment)  


import re
import os
import json
import torch
import uuid
import shutil
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple
from dataclasses import dataclass
from pathlib import Path

# 向量数据库相关库
import faiss
import langchain
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_community.embeddings import HuggingFaceEmbeddings


@dataclass
class Segment:
    """表示一个文档片段"""
    content: str
    metadata: dict
    index: int

@dataclass
class UploadResult:
    """表示文件上传和处理的结果"""
    success: bool
    message: str
    file_id: Optional[str] = None
    segments_count: Optional[int] = None
    vector_db_path: Optional[str] = None
    error: Optional[str] = None

@dataclass
class SearchResult:
    """表示检索结果"""
    content: str
    metadata: dict
    score: float
    source_file: str

class KnowledgeBaseSearcher:
    """知识库检索器，用于搜索已建立的向量数据库"""
    
    def __init__(self, vector_base_dir: str = "vector", model_path: str = "/bigdata/gpustack/cache/modelscope/hub/models/BAAI/bge-m3"):
        """
        初始化知识库检索器
        
        Args:
            vector_base_dir: 向量数据库基础目录
            model_path: 向量化模型路径
        """
        self.vector_base_dir = vector_base_dir
        self.model_path = model_path
        self.embeddings = None
        self.loaded_databases = {}  # 缓存已加载的数据库
        
    def _init_embeddings(self):
        """初始化嵌入模型"""
        if self.embeddings is None:
            model_kwargs = {
                'device': 'cuda' if torch.cuda.is_available() else 'cpu',
            }
            
            encode_kwargs = {
                'normalize_embeddings': True,
            }
            
            class BGEEmbeddings(HuggingFaceEmbeddings):
                def embed_query(self, text: str) -> List[float]:
                    text = f"查询: {text}"
                    return super().embed_query(text)
            
            self.embeddings = BGEEmbeddings(
                model_name=self.model_path,
                model_kwargs=model_kwargs,
                encode_kwargs=encode_kwargs,
            )
    
    def get_available_databases(self) -> List[Dict[str, Any]]:
        """获取可用的向量数据库列表"""
        databases = []
        
        if not os.path.exists(self.vector_base_dir):
            return databases
        
        # 读取处理记录
        records_file = os.path.join(self.vector_base_dir, "processing_records.json")
        if os.path.exists(records_file):
            try:
                with open(records_file, 'r', encoding='utf-8') as f:
                    records = json.load(f)
                
                for record in records:
                    if os.path.exists(record['vector_db_path']):
                        databases.append({
                            'file_id': record['file_id'],
                            'filename': record['original_filename'],
                            'segments_count': record['segments_count'],
                            'processed_at': record['processed_at'],
                            'vector_db_path': record['vector_db_path']
                        })
            except Exception as e:
                print(f"读取处理记录失败: {str(e)}")
        
        return databases
    
    def load_database(self, db_path: str) -> Optional[FAISS]:
        """加载指定的向量数据库"""
        try:
            if db_path in self.loaded_databases:
                return self.loaded_databases[db_path]
            
            self._init_embeddings()
            
            if not os.path.exists(db_path):
                print(f"数据库路径不存在: {db_path}")
                return None
            
            vectorstore = FAISS.load_local(db_path, self.embeddings, allow_dangerous_deserialization=True)
            self.loaded_databases[db_path] = vectorstore
            return vectorstore
            
        except Exception as e:
            print(f"加载数据库失败 {db_path}: {str(e)}")
            return None
    
    def search_single_database(self, query: str, db_path: str, top_k: int = 5) -> List[SearchResult]:
        """在单个数据库中搜索"""
        vectorstore = self.load_database(db_path)
        if vectorstore is None:
            return []
        
        try:
            # 执行相似性搜索
            docs_with_scores = vectorstore.similarity_search_with_score(query, k=top_k)
            
            results = []
            for doc, score in docs_with_scores:
                result = SearchResult(
                    content=doc.page_content,
                    metadata=doc.metadata,
                    score=float(score),
                    source_file=doc.metadata.get('file_name', 'Unknown')
                )
                results.append(result)
            
            return results
            
        except Exception as e:
            print(f"搜索失败: {str(e)}")
            return []
    
    def search_all_databases(self, query: str, top_k: int = 10) -> List[SearchResult]:
        """在所有可用数据库中搜索"""
        all_results = []
        databases = self.get_available_databases()
        
        if not databases:
            return []
        
        # 在每个数据库中搜索
        for db_info in databases:
            db_results = self.search_single_database(query, db_info['vector_db_path'], top_k=5)
            all_results.extend(db_results)
        
        # 按相似度分数排序（分数越低越相似）
        all_results.sort(key=lambda x: x.score)
        
        # 返回前top_k个结果
        return all_results[:top_k]
    
    def search_by_file_ids(self, query: str, file_ids: List[str], top_k: int = 5) -> List[SearchResult]:
        """在指定文件的数据库中搜索"""
        all_results = []
        databases = self.get_available_databases()
        
        # 过滤出指定文件ID的数据库
        target_databases = [db for db in databases if db['file_id'] in file_ids]
        
        for db_info in target_databases:
            db_results = self.search_single_database(query, db_info['vector_db_path'], top_k=top_k)
            all_results.extend(db_results)
        
        # 按相似度分数排序
        all_results.sort(key=lambda x: x.score)
        
        return all_results[:top_k]

class DocumentUploadManager:
    """文档上传管理器，处理文件上传、存储和索引"""
    
    def __init__(self, upload_dir: str = "uploads", vector_base_dir: str = "vector"):
        """
        初始化文档上传管理器
        
        Args:
            upload_dir: 上传文件存储目录
            vector_base_dir: 向量数据库存储基础目录
        """
        self.upload_dir = upload_dir
        self.vector_base_dir = vector_base_dir
        self.supported_extensions = {'.txt', '.docx', '.pdf'}
        
        # 创建必要的目录
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(vector_base_dir, exist_ok=True)
        
        # 初始化知识库检索器
        self.searcher = KnowledgeBaseSearcher(vector_base_dir)
    
    def validate_file(self, filename: str, file_size: int, max_size_mb: int = 50) -> Dict[str, Any]:
        """
        验证上传的文件
        
        Args:
            filename: 文件名
            file_size: 文件大小（字节）
            max_size_mb: 最大文件大小（MB）
        
        Returns:
            验证结果字典
        """
        # 检查文件扩展名
        ext = os.path.splitext(filename)[1].lower()
        if ext not in self.supported_extensions:
            return {
                "valid": False,
                "error": f"不支持的文件格式: {ext}。支持的格式: {', '.join(self.supported_extensions)}"
            }
        
        # 检查文件大小
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return {
                "valid": False,
                "error": f"文件大小超过限制。最大允许: {max_size_mb}MB，当前文件: {file_size / 1024 / 1024:.2f}MB"
            }
        
        return {"valid": True}
    
    async def save_uploaded_file(self, file_content: bytes, original_filename: str) -> Dict[str, Any]:
        """
        保存上传的文件
        
        Args:
            file_content: 文件内容
            original_filename: 原始文件名
        
        Returns:
            保存结果字典
        """
        try:
            # 生成唯一的文件ID
            file_id = str(uuid.uuid4())
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # 保持原始文件扩展名
            ext = os.path.splitext(original_filename)[1]
            safe_filename = f"{timestamp}_{file_id}{ext}"
            
            # 保存文件
            file_path = os.path.join(self.upload_dir, safe_filename)
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            return {
                "success": True,
                "file_id": file_id,
                "file_path": file_path,
                "safe_filename": safe_filename,
                "original_filename": original_filename
            }
        
        except Exception as e:
            return {
                "success": False,
                "error": f"保存文件失败: {str(e)}"
            }
    
    def process_uploaded_file(self, file_path: str, file_id: str, original_filename: str, 
                            delimiter: str = '@@', keep_delimiter: bool = False) -> UploadResult:
        """
        处理上传的文件，进行分割和向量化
        
        Args:
            file_path: 文件路径
            file_id: 文件ID
            original_filename: 原始文件名
            delimiter: 分割符
            keep_delimiter: 是否保留分割符
        
        Returns:
            处理结果
        """
        try:
            # 创建文档分割器
            segmenter = DocumentSegmenter(delimiter=delimiter, keep_delimiter=keep_delimiter)
            
            # 处理文件
            segments = segmenter.process_file(file_path)
            
            if not segments:
                return UploadResult(
                    success=False,
                    message="文件处理失败：未能提取到有效内容",
                    error="文件可能为空或格式不正确"
                )
            
            # 创建向量数据库存储目录
            vector_db_name = f"doc_{file_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            vector_db_path = os.path.join(self.vector_base_dir, vector_db_name)
            
            # 保存到向量数据库
            segmenter.save_to_vector_db(segments, vector_db_path)
            
            # 保存处理记录
            self._save_processing_record(file_id, original_filename, file_path, 
                                       vector_db_path, len(segments), delimiter)
            
            return UploadResult(
                success=True,
                message=f"文件处理成功！共分割出 {len(segments)} 个文本块并已存入向量数据库",
                file_id=file_id,
                segments_count=len(segments),
                vector_db_path=vector_db_path
            )
        
        except Exception as e:
            return UploadResult(
                success=False,
                message="文件处理失败",
                error=str(e)
            )
    
    def search_knowledge_base(self, query: str, file_ids: Optional[List[str]] = None, top_k: int = 5) -> List[SearchResult]:
        """
        搜索知识库
        
        Args:
            query: 查询文本
            file_ids: 指定搜索的文件ID列表，如果为None则搜索所有
            top_k: 返回结果数量
        
        Returns:
            搜索结果列表
        """
        if file_ids:
            return self.searcher.search_by_file_ids(query, file_ids, top_k)
        else:
            return self.searcher.search_all_databases(query, top_k)
    
    def _save_processing_record(self, file_id: str, original_filename: str, file_path: str,
                              vector_db_path: str, segments_count: int, delimiter: str):
        """保存处理记录到JSON文件"""
        record = {
            "file_id": file_id,
            "original_filename": original_filename,
            "file_path": file_path,
            "vector_db_path": vector_db_path,
            "segments_count": segments_count,
            "delimiter": delimiter,
            "processed_at": datetime.now().isoformat(),
            "status": "completed"
        }
        
        records_file = os.path.join(self.vector_base_dir, "processing_records.json")
        
        # 读取现有记录
        records = []
        if os.path.exists(records_file):
            try:
                with open(records_file, 'r', encoding='utf-8') as f:
                    records = json.load(f)
            except:
                records = []
        
        # 添加新记录
        records.append(record)
        
        # 保存记录
        with open(records_file, 'w', encoding='utf-8') as f:
            json.dump(records, f, ensure_ascii=False, indent=2)
    
    def get_processing_history(self) -> List[Dict[str, Any]]:
        """获取处理历史记录"""
        records_file = os.path.join(self.vector_base_dir, "processing_records.json")
        
        if not os.path.exists(records_file):
            return []
        
        try:
            with open(records_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    
    def delete_file_and_index(self, file_id: str) -> Dict[str, Any]:
        """删除文件和对应的向量索引"""
        try:
            records = self.get_processing_history()
            target_record = None
            
            # 查找对应记录
            for record in records:
                if record.get('file_id') == file_id:
                    target_record = record
                    break
            
            if not target_record:
                return {"success": False, "error": "未找到对应的文件记录"}
            
            # 删除原文件
            if os.path.exists(target_record['file_path']):
                os.remove(target_record['file_path'])
            
            # 删除向量数据库目录
            if os.path.exists(target_record['vector_db_path']):
                shutil.rmtree(target_record['vector_db_path'])
            
            # 从缓存中移除
            if target_record['vector_db_path'] in self.searcher.loaded_databases:
                del self.searcher.loaded_databases[target_record['vector_db_path']]
            
            # 更新记录
            updated_records = [r for r in records if r.get('file_id') != file_id]
            records_file = os.path.join(self.vector_base_dir, "processing_records.json")
            with open(records_file, 'w', encoding='utf-8') as f:
                json.dump(updated_records, f, ensure_ascii=False, indent=2)
            
            return {"success": True, "message": "文件和索引删除成功"}
        
        except Exception as e:
            return {"success": False, "error": f"删除失败: {str(e)}"}

class DocumentSegmenter:
    """文档内容分割器，支持按指定分隔符分割文档内容并存储到向量数据库"""
    
    def __init__(self, delimiter: str = '@@', keep_delimiter: bool = False):
        """
        初始化文档分割器
        
        Args:
            delimiter: 用于分割文档的分隔符，默认为 '@@'
            keep_delimiter: 是否在分割后的内容中保留分隔符，默认为 False
        """
        self.delimiter = delimiter
        self.keep_delimiter = keep_delimiter
    
    def read_file(self, file_path: str) -> str:
        """
        读取文档内容
        
        Args:
            file_path: 文档文件路径
        
        Returns:
            文档内容字符串
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 根据文件扩展名选择读取方式
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return self._read_text_file(file_path)
        elif ext == '.docx':
            return self._read_docx_file(file_path)
        elif ext == '.pdf':
            return self._read_pdf_file(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _read_text_file(self, file_path: str) -> str:
        """读取文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_docx_file(self, file_path: str) -> str:
        """读取docx文件"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("读取docx文件需要安装python-docx库: pip install python-docx")
        
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    
    def _read_pdf_file(self, file_path: str) -> str:
        """读取PDF文件"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("读取PDF文件需要安装pdfplumber库: pip install pdfplumber")
        
        with pdfplumber.open(file_path) as pdf:
            return '\n'.join([page.extract_text() or '' for page in pdf.pages])
    
    def split_content(self, content: str, metadata: Optional[dict] = None) -> List[Segment]:
        """
        按指定分隔符分割文档内容并生成带元数据的片段
        
        Args:
            content: 文档内容字符串
            metadata: 可选的元数据字典，将添加到每个分割段中
        
        Returns:
            分割后的Segment对象列表
        """
        if metadata is None:
            metadata = {}
        
        if not self.keep_delimiter:
            # 不保留分隔符，直接分割
            raw_segments = [section.strip() for section in content.split(self.delimiter) if section.strip()]
        else:
            # 保留分隔符，使用正则表达式分割
            pattern = re.compile(f'({re.escape(self.delimiter)})')
            parts = pattern.split(content)
            
            raw_segments = []
            current_segment = ''
            
            for part in parts:
                if part == self.delimiter:
                    if current_segment.strip():
                        raw_segments.append(current_segment.strip())
                    current_segment = self.delimiter
                else:
                    current_segment += part
            
            if current_segment.strip():
                raw_segments.append(current_segment.strip())
        
        # 生成带元数据的Segment对象
        segments = []
        for i, segment in enumerate(raw_segments):
            
            # 创建元数据
            segment_metadata = {
                "segment_index": i,
                "total_segments": len(raw_segments),
                **metadata
            }
            
            segments.append(Segment(
                content=segment,
                metadata=segment_metadata,
                index=i
            ))
        
        return segments
    

    
    def process_file(self, file_path: str) -> List[Segment]:
        """
        读取文件并按分隔符分割内容
        
        Args:
            file_path: 文档文件路径
        
        Returns:
            分割后的Segment对象列表
        """
        content = self.read_file(file_path)
        
        # 创建基础元数据
        base_metadata = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "creation_time": os.path.getctime(file_path)
        }
        
        return self.split_content(content, base_metadata)
    
    def save_to_vector_db(self, segments: List[Segment], save_dir: str, model_path: str = "/bigdata/gpustack/cache/modelscope/hub/models/BAAI/bge-m3") -> None:
        """
        将分割的文本块存入向量数据库并保存到指定目录
        """
        os.makedirs(save_dir, exist_ok=True)
        
        langchain_docs = [
            Document(
                page_content=segment.content,
                metadata=segment.metadata
            )
            for segment in segments
        ]
        
        model_kwargs = {
            'device': 'cuda' if torch.cuda.is_available() else 'cpu',
        }
        
        # 仅保留normalize_embeddings，移除show_progress_bar
        encode_kwargs = {
            'normalize_embeddings': True,
        }
        
        class BGEEmbeddings(HuggingFaceEmbeddings):
            def embed_query(self, text: str) -> List[float]:
                text = f"查询: {text}"
                return super().embed_query(text)
        
        embeddings = BGEEmbeddings(
            model_name=model_path,
            model_kwargs=model_kwargs,
            encode_kwargs=encode_kwargs,
        )
        
        vectorstore = FAISS.from_documents(langchain_docs, embeddings)
        vectorstore.save_local(save_dir)
        
        # 保存元数据索引
        metadata_path = os.path.join(save_dir, "metadata_index.json")
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump([segment.metadata for segment in segments], f, ensure_ascii=False, indent=2)
        
        print(f"成功将 {len(segments)} 个文本块存入向量数据库并保存到 {save_dir}")

# 示例用法
if __name__ == "__main__":
    # 创建分割器，使用'@@'作为分隔符，不保留分隔符
    segmenter = DocumentSegmenter(delimiter='@@', keep_delimiter=False)
    
    # 处理文档文件
    file_path = '/bigdata/lihuanjia/ocr/data/ERDOS_零售管理项目_培训手册_WEBPOS_20250210V1.docx'  # 替换为实际文件路径
    segments = segmenter.process_file(file_path)
    
    # 保存到向量数据库，使用BGE-m3模型
    save_dir = '/bigdata/lihuanjia/ocr/vector/erdos_webpos_training'
    segmenter.save_to_vector_db(segments, save_dir)
    
    # 输出分割结果摘要
    print(f"成功从文档中分割出 {len(segments)} 个内容块并保存到向量数据库")
    for i, segment in enumerate(segments, 1):  
        title = segment.metadata.get('segment_title', f"无标题片段 #{i}")
        print(f"\n=== 第 {i} 个内容块: {title} ===")
        print(segment.content[:200] + "..." if len(segment.content) > 200 else segment.content)  